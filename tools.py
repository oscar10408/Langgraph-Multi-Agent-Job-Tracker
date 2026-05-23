"""
tools.py
--------
Data layer utility functions.

Each function here has a single responsibility:
- Only handles data reading or writing
- Contains no LLM logic
- Returns plain text for easy LLM consumption

To switch to a real API (Notion, Airtable, etc.), only this file needs to change.
"""
import base64
import io
import json
import os
import re
import shutil
from datetime import date

import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from langchain_core.messages import HumanMessage
from langchain_groq import ChatGroq
from openpyxl import load_workbook
from playwright.sync_api import sync_playwright

# ── Default (owner) paths ───────────────────────────────────
_DEFAULT_EXCEL_PATH = os.path.join(os.path.dirname(__file__), "data", "Record_of_job_AI.xlsx")
_DEFAULT_PROFILE_PATH = os.path.join(os.path.dirname(__file__), "data", "profile.json")

# ---------------------------------------------------------------------------
# Session-aware helpers
# These functions check st.session_state for a custom upload first,
# and fall back to the owner's default files if nothing is uploaded.
# ---------------------------------------------------------------------------

def _get_excel_path() -> str:
    """Return the active Excel path (custom upload takes priority)."""
    return _DEFAULT_EXCEL_PATH


def _load_excel() -> tuple:
    """
    Load Excel and return (workbook, worksheet).
    Uses in-memory bytes from st.session_state if a custom file was uploaded,
    otherwise reads from the default path on disk.
    """
    custom_bytes = st.session_state.get("custom_excel_bytes")
    if custom_bytes:
        wb = load_workbook(io.BytesIO(custom_bytes))
    else:
        wb = load_workbook(_DEFAULT_EXCEL_PATH)
    ws = wb.active
    return wb, ws


def _save_excel(wb) -> None:
    """
    Save workbook back to the correct destination.
    For custom sessions: update the in-memory bytes in st.session_state.
    For the default session: write directly to disk (with a tmp swap).
    """
    custom_bytes = st.session_state.get("custom_excel_bytes")
    if custom_bytes is not None:
        buf = io.BytesIO()
        wb.save(buf)
        st.session_state["custom_excel_bytes"] = buf.getvalue()
    else:
        tmp = _DEFAULT_EXCEL_PATH + ".tmp"
        wb.save(tmp)
        shutil.move(tmp, _DEFAULT_EXCEL_PATH)


def _get_headers(ws) -> list:
    return [cell.value for cell in ws[1]]


def get_all_applications(limit: int = None) -> str:
    """
    Returns a summary of all job application records.
    limit: Only return the latest N records. None means return all.
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    rows = list(ws.iter_rows(min_row=2, values_only=True))

    if not rows:
        return "No job applications found."

    # Filter out empty rows
    valid_rows = [
        (i + 2, dict(zip(headers, row)))
        for i, row in enumerate(rows)
        if any(row)
    ]

    # If limit is set, return only the latest N records
    if limit:
        valid_rows = valid_rows[-limit:]

    lines = [f"=== Job Applications (showing {len(valid_rows)} records) ==="]
    for _, d in valid_rows:
        status = d.get("Status") or "unknown"
        status_emoji = {
            "applied": "📤", "interviewed": "🎯",
            "rejected": "❌", "offered": "🎉", "wishlist": "⭐"
        }.get(str(status).lower(), "❓")

        lines.append(
            f"{status_emoji} {d.get('Company', 'N/A')} - {d.get('Position', 'N/A')}\n"
            f"   Status: {status} | Applied: {d.get('Applied on', 'N/A')}\n"
            f"   Link: {str(d.get('Application Link', ''))}"
        )
    return "\n\n".join(lines)


def get_pending_followups() -> str:
    """Returns a list of applications still in applied status."""
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    today = date.today().isoformat()

    pending = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        d = dict(zip(headers, row))
        status = str(d.get("Status") or "").lower()
        applied_on = str(d.get("Applied on") or "")
        if status == "applied" and applied_on and applied_on <= today:
            pending.append(d)

    if not pending:
        return "No follow-ups needed right now. You're on top of things!"

    lines = ["=== Pending Follow-ups ==="]
    for d in pending:
        lines.append(
            f"⏰ {d.get('Company')} - {d.get('Position')}\n"
            f"   Applied: {d.get('Applied on')}"
        )
    return "\n\n".join(lines)


def add_application(company: str, role: str, notes: str = "") -> str:
    """Adds a new job application record to Excel."""
    wb, ws = _load_excel()
    today = date.today().isoformat()
    ws.append([company, role, notes, "", today, "applied"])
    _save_excel(wb)

    return f"✅ Added: {company} - {role}"


def get_status_summary() -> str:
    """Returns a statistical summary of application statuses."""
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    counts = {}
    total = 0

    for row in ws.iter_rows(min_row=2, values_only=True):
        d = dict(zip(headers, row))
        if not d.get("Company"):
            continue
        status = str(d.get("Status") or "unknown").lower()
        counts[status] = counts.get(status, 0) + 1
        total += 1

    lines = ["=== Application Summary ==="]
    for status, count in counts.items():
        lines.append(f"  {status}: {count}")
    lines.append(f"\n  Total: {total}")
    return "\n".join(lines)


def update_application(company: str = None, role: str = None,
                        notes: str = None, status: str = None,
                        job_id: int = None) -> str:
    """
    Updates fields for a specific application record, identified by row index or company name.
    job_id corresponds to the Nth data row in Excel (starting from 1).
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    col_map = {name: idx + 1 for idx, name in enumerate(headers)}

    target_row = None

    if job_id:
        target_row = job_id + 1  # +1 because row 1 is the header
    else:
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            d = dict(zip(headers, row))
            if company and str(d.get("Company") or "").lower() == company.lower():
                target_row = i
                break

    if not target_row:
        return "❌ Job not found."

    if company and "Company" in col_map:
        ws.cell(row=target_row, column=col_map["Company"]).value = company
    if role and "Position" in col_map:
        ws.cell(row=target_row, column=col_map["Position"]).value = role
    if notes and "JD" in col_map:
        ws.cell(row=target_row, column=col_map["JD"]).value = notes
    if status and "Status" in col_map:
        ws.cell(row=target_row, column=col_map["Status"]).value = status

    _save_excel(wb)

    row_data = dict(zip(headers, list(ws.iter_rows(min_row=target_row,
                    max_row=target_row, values_only=True))[0]))
    return f"✅ Updated: {row_data.get('Company')} - {row_data.get('Position')}"


def scrape_job_url(url: str) -> str:
    """
    Scrape job description from a URL.
    For Workday sites: use requests (meta tags contain full JD).
    For others: fall back to Playwright.
    """

    # ── Workday: fetch JD directly from meta tags via requests ──
    if "myworkdayjobs.com" in url:
        try:
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                )
            }
            resp = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(resp.text, "html.parser")

            # Workday embeds the full JD in meta[name="description"] or meta[property="og:description"]
            content = ""
            for attr in [("name", "description"), ("property", "og:description")]:
                tag = soup.find("meta", {attr[0]: attr[1]})
                if tag and tag.get("content"):
                    content = tag["content"]
                    break

            if content:
                print(f"[Scrape] ✅ Workday meta fetched {len(content)} chars from {url}")
                return content
        except Exception as e:
            print(f"[Scrape] ❌ Workday requests failed: {e}")
            # fall through to Playwright

    # ── Other sites: fall back to Playwright ──
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-dev-shm-usage"]
            )
            page = browser.new_page()
            page.goto(url, timeout=20000)
            page.wait_for_load_state("networkidle", timeout=20000)
            page.wait_for_timeout(2000)
            html = page.content()
            browser.close()

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        content = soup.get_text(separator="\n", strip=True)
        print(f"[Scrape] ✅ Playwright fetched {len(content)} chars from {url}")
        return content
    except Exception as e:
        print(f"[Scrape] ❌ Playwright failed: {e}")
        return ""

def get_incomplete_rows() -> list[dict]:
    """
    Reads Excel and finds rows that have an Application Link but are missing Company, Position, or JD.
    Returns a list of incomplete row data.
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    incomplete = []

    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        row_dict = dict(zip(headers, row))
        url = row_dict.get("Application Link", "")
        company = row_dict.get("Company", "")
        position = row_dict.get("Position", "")
        jd = row_dict.get("JD", "")

        if url and (not company or not position or not jd):
            incomplete.append({
                "row_index": i,
                "url": url,
                "company": company,
                "position": position,
                "jd": jd,
            })

    return incomplete


def update_excel_row(row_index: int, company: str = None,
                      position: str = None, jd: str = None) -> None:
    """
    Updates Company, Position, and JD fields for a specific Excel row.
    Only updates fields with non-None values; None means skip.
    """
    wb, _ = _load_excel()
    ws = wb.active

    headers = [cell.value for cell in ws[1]]
    col_map = {name: idx + 1 for idx, name in enumerate(headers)}

    if company and "Company" in col_map:
        ws.cell(row=row_index, column=col_map["Company"]).value = company
    if position and "Position" in col_map:
        ws.cell(row=row_index, column=col_map["Position"]).value = position
    if jd and "JD" in col_map:
        ws.cell(row=row_index, column=col_map["JD"]).value = jd

    _save_excel(wb)



def load_profile() -> dict:
    """
    Load user profile.
    Returns the custom profile from st.session_state if one has been uploaded,
    otherwise falls back to the owner's default profile.json on disk.
    """
    custom = st.session_state.get("custom_profile")
    if custom:
        return custom
    with open(_DEFAULT_PROFILE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_cover_letter(company: str, position: str, content: str) -> str:
    """
    Create a folder named after the company and save the cover letter as a Word file.
    """
    # Sanitize company and position names by removing characters invalid in folder/file names
    safe_company = re.sub(r'[\\/*?:"<>|]', "", company).strip()
    safe_position = re.sub(r'[\\/*?:"<>|]', "", position).strip()

    # Build output directory path
    output_dir = os.path.join(os.path.dirname(__file__), "cover_letters", safe_company)
    os.makedirs(output_dir, exist_ok=True)

    # Create Word document
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.add_heading(f"Cover Letter — {safe_position}", level=1)
    doc.add_paragraph(f"Company: {safe_company}")
    doc.add_paragraph(f"Position: {safe_position}")
    doc.add_paragraph("")  # blank line

    for para in content.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Save Word file
    filename = f"Cover_Letter_{safe_position.replace(' ', '_')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    # Also export as PDF
    try:
        pdf_path = filepath.replace(".docx", ".pdf")
        convert(filepath, pdf_path)
        print(f"[CoverLetter] PDF saved to: {pdf_path}")
    except Exception as e:
        print(f"[CoverLetter] ⚠️ PDF conversion failed: {e}")

    return filepath

GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]

def _get_gmail_service():
    """
    Build a Gmail API service client.
    Priority:
      1. Custom token uploaded by the user in this session (st.session_state)
      2. Owner's token stored in st.secrets (default)
    """
    custom_token = st.session_state.get("custom_gmail_token")
    if custom_token:
        token_dict = dict(custom_token)
    else:
        token_dict = dict(st.secrets["gmail_token"])

    # scopes may be a toml array or a plain list; normalise to list of str
    if "scopes" in token_dict:
        token_dict["scopes"] = list(token_dict["scopes"])

    creds = Credentials.from_authorized_user_info(token_dict, GMAIL_SCOPES)

    # Refresh automatically if token is expired (no browser needed)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        # ⚠️ Note: refreshed token only lives in memory; will be re-refreshed after reboot
        # This is acceptable since the refresh_token remains valid long-term

    return build("gmail", "v1", credentials=creds)

def scan_emails_for_status(max_results: int = 50) -> str:
    """
    Scans recent emails to find job application status updates.
    Uses LLM to extract company name and status, reducing false positives.
    """

    _llm = ChatGroq(model="meta-llama/llama-4-scout-17b-16e-instruct")

    service = _get_gmail_service()

    keywords = "subject:(interview OR application OR offer OR unfortunately OR position OR opportunity)"
    results = service.users().messages().list(
        userId="me", q=keywords, maxResults=max_results
    ).execute()

    messages = results.get("messages", [])
    if not messages:
        return "📭 No job-related emails found."

    wb, ws = _load_excel()
    headers = _get_headers(ws)
    col_map = {name: idx + 1 for idx, name in enumerate(headers)}

    # Build company name index for faster lookups
    company_index = {}
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        d = dict(zip(headers, row))
        company = str(d.get("Company") or "").strip()
        if company:
            company_index[company.lower()] = (i, d)

    updates = []

    for msg in messages:
        msg_data = service.users().messages().get(
            userId="me", id=msg["id"], format="full"
        ).execute()

        headers_list = msg_data["payload"].get("headers", [])
        subject = next((h["value"] for h in headers_list if h["name"] == "Subject"), "")
        sender = next((h["value"] for h in headers_list if h["name"] == "From"), "")

        # Extract email body text
        body = ""
        payload = msg_data["payload"]
        if "parts" in payload:
            for part in payload["parts"]:
                if part["mimeType"] == "text/plain":
                    data = part["body"].get("data", "")
                    body = base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")
                    break
        elif "body" in payload:
            data = payload["body"].get("data", "")
            body = base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")

        # Use LLM to extract both company name and application status
        classify_prompt = f"""You are analyzing a job application email for Oscar Shih.

        Email:
        Subject: {subject}
        From: {sender}
        Body: {body}

        Answer these two questions:
        1. What is the HIRING company's name?
        - NOT email platforms: Workday, Dayforce, Indeed, LinkedIn, Microsoft Teams, Handshake
        - NOT if this is a job recommendation or advertisement (Handshake recommending jobs, recruiter cold outreach)
        - Only if this email is about an application Oscar Shih already submitted

        2. What is the application status?

        Return ONLY in this exact format:
        company: <hiring company name, or empty if not applicable>
        status: <interviewed / offered / rejected / applied / ignore>

        Status rules:
        - interviewed: scheduling or confirming an interview
        - offered: job offer or candidate is selected
        - rejected: not selected, moving forward with other candidates
        - applied: confirming application was received
        - ignore: job recommendations, advertisements, cold outreach, newsletters, or unrelated emails"""

        classify_response = _llm.invoke([HumanMessage(content=classify_prompt)])
        classify_text = classify_response.content.strip()

        detected_company, new_status = None, None
        for line in classify_text.split("\n"):
            if line.startswith("company:"):
                val = line.split(":", 1)[1].strip()
                detected_company = val if val else None
            elif line.startswith("status:"):
                val = line.split(":", 1)[1].strip().lower()
                new_status = val if val in ["interviewed", "offered", "rejected"] else None

        # Skip if ignored or no usable information extracted
        if not detected_company or not new_status:
            print(f"[EmailScan] Skipped: {subject}")
            continue

        print(f"[EmailScan] Detected: {detected_company} → {new_status} | {subject[:50]}")

        # Find the closest matching company in the index
        target_row = None
        for excel_company_lower, (i, d) in company_index.items():
            if (detected_company.lower() in excel_company_lower or
                excel_company_lower in detected_company.lower()):
                target_row = (i, d)
                break

        if not target_row:
            print(f"[EmailScan] ⚠️ '{detected_company}' not found in Excel, skipping.")
            continue

        i, d = target_row
        current_status = str(d.get("Status") or "").lower()
        if current_status != new_status:
            ws.cell(row=i, column=col_map["Status"]).value = new_status
            updates.append(f"  {detected_company} → {new_status} (from: {subject[:50]})")

    if updates:
        _save_excel(wb)
        return f"✅ Updated {len(updates)} applications:\n" + "\n".join(updates)
    else:
        return f"📭 Scanned {len(messages)} emails, no status updates detected."



def save_interview_prep(company: str, position: str, content: str) -> str:
    """
    Creates a company-named folder and saves the interview prep as a Word document.
    """
    safe_company = re.sub(r'[\\/*?:"<>|]', "", company).strip()
    safe_position = re.sub(r'[\\/*?:"<>|]', "", position).strip()

    output_dir = os.path.join(os.path.dirname(__file__), "interview_prep", safe_company)
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.add_heading(f"Interview Prep — {safe_position} @ {safe_company}", level=1)
    doc.add_paragraph("")

    for line in content.split("\n"):
        if line.strip().startswith("##"):
            doc.add_heading(line.strip().replace("##", "").strip(), level=2)
        elif line.strip().startswith("#"):
            doc.add_heading(line.strip().replace("#", "").strip(), level=1)
        elif line.strip():
            doc.add_paragraph(line.strip())

    filename = f"Interview_Prep_{safe_position.replace(' ', '_')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"[InterviewPrep] Word saved to: {filepath}")

    try:

        pdf_path = filepath.replace(".docx", ".pdf")
        convert(filepath, pdf_path)
        print(f"[InterviewPrep] PDF saved to: {pdf_path}")
    except Exception as e:
        print(f"[InterviewPrep] ⚠️ PDF conversion failed: {e}")

    return filepath

def get_jd_from_excel(company: str = None, position: str = None) -> dict:
    """
    Finds the matching JD and URL from Excel by company or position name.
    Returns {jd, url, company, position, row_index} or None if not found.
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)

    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        d = dict(zip(headers, row))
        excel_company = str(d.get("Company") or "").lower()
        excel_position = str(d.get("Position") or "").lower()

        match = False
        if company and company.lower() in excel_company:
            match = True
        if position and position.lower() in excel_position:
            match = True

        if match:
            return {
                "jd": str(d.get("JD") or "").strip(),
                "url": str(d.get("Application Link") or "").strip(),
                "company": d.get("Company", ""),
                "position": d.get("Position", ""),
                "row_index": i
            }
    return None

def save_match_score_to_excel(row_index: int, score: str) -> None:
    """
    Writes the match score back to the Match Score column of the corresponding Excel row.
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    col_map = {name: idx + 1 for idx, name in enumerate(headers)}

    if "Match Score" not in col_map:
        print("[JobMatch] ⚠️ 'Match Score' column not found in Excel")
        return

    ws.cell(row=row_index, column=col_map["Match Score"]).value = score

    _save_excel(wb)
    print(f"[JobMatch] ✅ Score {score} saved to Excel row {row_index}")


def get_latest_jobs_with_jd(limit: int = 5) -> list[dict]:
    """
    Returns the latest N job records that have either a JD or an Application Link.
    """
    wb, ws = _load_excel()
    headers = _get_headers(ws)
    rows = list(ws.iter_rows(min_row=2, values_only=True))

    valid_rows = []
    for i, row in enumerate(rows, start=2):
        d = dict(zip(headers, row))
        if not d.get("Company"):
            continue
        if d.get("JD") or d.get("Application Link"):
            valid_rows.append({"row_index": i, **d})

    # Return only the latest N records
    return valid_rows[-limit:]
