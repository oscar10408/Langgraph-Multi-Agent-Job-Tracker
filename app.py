"""
app.py
------
Streamlit Web UI for Job Application Tracker Agent.
Phase 1: Dashboard + Applications Table
Phase 2: AI Tools (Cover Letter, Interview Prep, Job Match)
Phase 3: Chat bar (Agent integration)
"""
from dotenv import load_dotenv
load_dotenv()

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from openpyxl import load_workbook
from datetime import date
import os
import subprocess
import sys

@st.cache_resource
def install_playwright():
    subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"], check=True)

install_playwright()

if "GROQ_API_KEY" in st.secrets:
    os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]
    
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# ── Page config ────────────────────────────────────────────
st.set_page_config(
    page_title="Job Tracker",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── Custom CSS ─────────────────────────────────────────────
st.markdown("""
<style>
    /* Hide Streamlit default elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* Main background */
    .stApp { background-color: #f8f7f4; }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #ffffff;
        border-right: 1px solid #e8e6e0;
    }

    /* Metric cards */
    [data-testid="stMetric"] {
        background: #ffffff;
        border: 1px solid #e8e6e0;
        border-radius: 10px;
        padding: 1rem 1.25rem;
    }

    /* Buttons */
    .stButton > button {
        background: #1a1a1a;
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 500;
        padding: 0.5rem 1.25rem;
    }
    .stButton > button:hover {
        background: #333333;
        color: white;
    }

    /* Status badges */
    .badge {
        display: inline-block;
        padding: 3px 10px;
        border-radius: 20px;
        font-size: 12px;
        font-weight: 500;
    }
    .badge-interviewed { background: #dbeafe; color: #1d4ed8; }
    .badge-applied { background: #f3f4f6; color: #4b5563; }
    .badge-offered { background: #dcfce7; color: #166534; }
    .badge-rejected { background: #fee2e2; color: #991b1b; }
    .badge-unknown { background: #f3f4f6; color: #6b7280; }
    .badge-wishlist { background: #fef9c3; color: #854d0e; }
    .badge-followup { background: #f3e8ff; color: #7e22ce; }

    /* Section headers */
    .section-header {
        font-size: 16px;
        font-weight: 600;
        color: #1a1a1a;
        margin-bottom: 0.75rem;
        margin-top: 1.5rem;
    }

    /* AI tool cards */
    .ai-card {
        background: white;
        border: 1px solid #e8e6e0;
        border-radius: 12px;
        padding: 1.25rem;
        cursor: pointer;
        transition: border-color 0.15s;
    }
    .ai-card:hover { border-color: #1a1a1a; }
    .ai-card-icon { font-size: 24px; margin-bottom: 8px; }
    .ai-card-title { font-size: 14px; font-weight: 600; color: #1a1a1a; }
    .ai-card-desc { font-size: 12px; color: #6b7280; margin-top: 4px; }

    /* Chat messages */
    .chat-msg-user {
        background: #1a1a1a;
        color: white;
        padding: 10px 14px;
        border-radius: 12px 12px 2px 12px;
        margin: 8px 0;
        margin-left: 20%;
        font-size: 14px;
    }
    .chat-msg-agent {
        background: white;
        border: 1px solid #e8e6e0;
        color: #1a1a1a;
        padding: 10px 14px;
        border-radius: 12px 12px 12px 2px;
        margin: 8px 0;
        margin-right: 20%;
        font-size: 14px;
    }

    /* Dataframe styling */
    .stDataFrame { border-radius: 10px; overflow: hidden; }

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: transparent;
    }
    .stTabs [data-baseweb="tab"] {
        background: white;
        border: 1px solid #e8e6e0;
        border-radius: 8px;
        color: #6b7280;
        font-size: 13px;
    }
    .stTabs [aria-selected="true"] {
        background: #1a1a1a !important;
        color: white !important;
        border-color: #1a1a1a !important;
    }
</style>
""", unsafe_allow_html=True)


# ── Data loading ───────────────────────────────────────────
EXCEL_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "Record_of_job_AI.xlsx")

@st.cache_data(ttl=30)
def load_data():
    try:
        df = pd.read_excel(EXCEL_PATH)
        df.columns = df.columns.str.strip()
        df = df[df["Company"].notna() & (df["Company"] != "")]

        # Ensure Stage column exists
        if "Stage" not in df.columns:
            df["Stage"] = ""

        df["Status"] = df["Status"].fillna("").str.strip()
        df["Stage"]  = df["Stage"].fillna("").str.strip()

        # Resolve final display status from Status + Stage (Stage as fallback)
        def _resolve(row):
            s = _normalize_status(row["Status"])
            if s != "Unknown":
                return s
            # Fall back to last segment of Stage
            stage = str(row["Stage"]).strip()
            if stage and stage.lower() not in ("none", ""):
                last = stage.split("\u2192")[-1].strip()
                s2 = _normalize_status(last)
                if s2 != "Unknown":
                    return s2
            return "Applied"  # blank status = still in progress

        df["Status"] = df.apply(_resolve, axis=1)
        df["Applied on"] = df["Applied on"].fillna("Unknown")
        return df
    except Exception as e:
        st.error(f"Error loading data: {e}")
        return pd.DataFrame()

def _normalize_status(raw: str) -> str:
        """Map free-form Status values to standard labels for display."""
        s = str(raw or "").lower().strip()
        if not s or s == "none":
            return "Unknown"
        if any(x in s for x in ["reject", "unfortunately", "not moving"]):
            return "Rejected"
        if any(x in s for x in ["offer", "selected", "hired"]):
            return "Offered"
        if any(x in s for x in ["interview", "screening", "1st", "2nd", "final", "phone", "onsite"]):
            return "Interviewed"
        if any(x in s for x in ["follow-up", "questionaire", "questionnaire", "assessment", "task"]):
            return "Follow-up Q"
        if s in ["applied", "pending"]:
            return "Applied"
        if s == "wishlist":
            return "Wishlist"
        return "Unknown"


def get_status_counts(df):
    """df["Status"] is already normalized by load_data — just count directly."""
    counts = df["Status"].value_counts().to_dict()
    return {
        "total":       len(df),
        "interviewed": counts.get("Interviewed", 0),
        "offered":     counts.get("Offered", 0),
        "rejected":    counts.get("Rejected", 0),
        "applied":     counts.get("Applied", 0),
        "follow_up":   counts.get("Follow-up Q", 0),
    }


def status_badge(status):
    """Status is already normalized — map directly to badge CSS."""
    badge_map = {
        "Interviewed": "badge-interviewed",
        "Offered":     "badge-offered",
        "Rejected":    "badge-rejected",
        "Applied":     "badge-applied",
        "Follow-up Q": "badge-followup",
        "Wishlist":    "badge-wishlist",
        "Unknown":     "badge-unknown",
    }
    label = str(status).strip()
    css = badge_map.get(label, "badge-unknown")
    return f'<span class="badge {css}">{label}</span>'


# ── Sidebar ────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 💼 Job Tracker")

    # Show whose profile is active
    using_custom = st.session_state.get("using_custom_profile", False)
    if using_custom:
        active_name = st.session_state.get("custom_profile", {}).get("name", "Custom user")
        st.markdown(f"**{active_name}**")
    else:
        from tools import load_profile
        try:
            _owner = load_profile()
            st.markdown(f"**{_owner.get('name', 'Owner')}** · {_owner.get('location', '')}")
        except Exception:
            st.markdown("**Job Tracker**")

    st.divider()

    page = st.radio(
        "Navigation",
        ["📊 Dashboard", "📋 Applications", "📧 Email Scanner", "📝 Cover Letter", "🎯 Interview Prep", "🔍 Job Match", "💬 Chat", "⚙️ Setup"],
        label_visibility="collapsed"
    )

    st.divider()

    if st.button("🔄 Refresh data", use_container_width=True):
        st.cache_data.clear()
        st.rerun()


# ── Load data ──────────────────────────────────────────────
df = load_data()
stats = get_status_counts(df) if not df.empty else {}


# ══════════════════════════════════════════════════════════
# PAGE: DASHBOARD
# ══════════════════════════════════════════════════════════
if page == "📊 Dashboard":
    st.markdown("## Dashboard")
    st.markdown(f"*Last updated: {date.today().strftime('%B %d, %Y')}*")

    # Stats row
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Total Applied", stats.get("total", 0))
    with col2:
        st.metric("Interviews", stats.get("interviewed", 0))
    with col3:
        rate = round(stats.get("interviewed", 0) / max(stats.get("total", 1), 1) * 100, 1)
        st.metric("Interview Rate", f"{rate}%")
    with col4:
        st.metric("Offers", stats.get("offered", 0))
    with col5:
        st.metric("Rejected", stats.get("rejected", 0))

    st.markdown("---")

    # ── Build status counts (shared by both charts) ──────────
    STATUS_COLORS = {
        "Interviewed": "#1d4ed8",
        "Applied":     "#4b5563",
        "Rejected":    "#991b1b",
        "Offered":     "#166534",
        "Follow-up Q": "#7e22ce",
        "Wishlist":    "#854d0e",
        "Unknown":     "#6b7280",
    }

    # df["Status"] already normalized in load_data — count directly
    status_counts = {}
    for label in ["Interviewed", "Applied", "Rejected", "Offered", "Follow-up Q", "Wishlist", "Unknown"]:
        count = int((df["Status"] == label).sum())
        if count > 0:
            status_counts[label] = count

    ordered_labels = list(status_counts.keys())
    ordered_colors = [STATUS_COLORS.get(label, "#6b7280") for label in ordered_labels]

    shared_layout = dict(
        plot_bgcolor="white",
        paper_bgcolor="white",
        margin=dict(l=0, r=0, t=10, b=0),
        height=300,
        font=dict(size=13),
    )

    tab_status, tab_funnel = st.tabs(["📊 Status breakdown", "🔀 Application funnel"])

    with tab_status:
        col_left, col_right = st.columns([1.5, 1])

        with col_left:
            st.markdown('<div class="section-header">Application status breakdown</div>', unsafe_allow_html=True)
            if status_counts:
                fig = px.bar(
                    x=ordered_labels,
                    y=list(status_counts.values()),
                    color=ordered_labels,
                    color_discrete_map=STATUS_COLORS,
                    labels={"x": "", "y": "Count"},
                )
                fig.update_layout(showlegend=False, **shared_layout)
                fig.update_traces(marker_line_width=0)
                st.plotly_chart(fig, use_container_width=True)

        with col_right:
            st.markdown('<div class="section-header">Status distribution</div>', unsafe_allow_html=True)
            if status_counts:
                fig2 = go.Figure(data=[go.Pie(
                    labels=ordered_labels,
                    values=list(status_counts.values()),
                    hole=0.55,
                    marker_colors=ordered_colors,
                )])
                fig2.update_layout(
                    showlegend=True,
                    legend=dict(font=dict(size=11)),
                    **shared_layout,
                )
                fig2.update_traces(textinfo="none")
                st.plotly_chart(fig2, use_container_width=True)

    with tab_funnel:
        st.markdown('<div class="section-header">Application funnel</div>', unsafe_allow_html=True)
        stage_data = df[df["Stage"].notna() & (df["Stage"].astype(str).str.strip() != "")]
        if len(stage_data) < 5:
            st.info(
                f"Not enough stage data yet ({len(stage_data)} records). "
                "The funnel chart will appear once more applications have stage history. "
                "Stage is automatically recorded whenever you update an application status."
            )
        else:
            # Parse stage paths into Sankey links
            from collections import defaultdict
            link_counts = defaultdict(int)
            for stage_str in stage_data["Stage"].astype(str):
                nodes = [s.strip() for s in stage_str.split("→")]
                for i in range(len(nodes) - 1):
                    link_counts[(nodes[i], nodes[i+1])] += 1

            all_nodes = list(dict.fromkeys(
                n for pair in link_counts for n in pair
            ))
            node_idx = {n: i for i, n in enumerate(all_nodes)}

            node_colors = {
                "interviewed":  "#1d4ed8",
                "offered":      "#166534",
                "rejected":     "#991b1b",
                "follow-up q":  "#7e22ce",
            }
            n_colors = [node_colors.get(n.lower(), "#4b5563") for n in all_nodes]

            fig3 = go.Figure(go.Sankey(
                node=dict(
                    label=all_nodes,
                    color=n_colors,
                    pad=20,
                    thickness=20,
                ),
                link=dict(
                    source=[node_idx[s] for s, _ in link_counts],
                    target=[node_idx[t] for _, t in link_counts],
                    value=list(link_counts.values()),
                    color="rgba(180,180,180,0.3)",
                ),
            ))
            fig3.update_layout(
                paper_bgcolor="white",
                margin=dict(l=0, r=0, t=10, b=0),
                height=320,
                font=dict(size=13),
            )
            st.plotly_chart(fig3, use_container_width=True)

    # Recent applications
    st.markdown('<div class="section-header">Recent applications</div>', unsafe_allow_html=True)

    recent = df.tail(10)[["Company", "Position", "Applied on", "Status"]].copy()
    recent = recent.iloc[::-1]

    for _, row in recent.iterrows():
        col1, col2, col3, col4 = st.columns([2, 2.5, 1.5, 1.5])
        with col1:
            st.markdown(f"**{row['Company']}**")
        with col2:
            pos = str(row['Position'])
            st.markdown(f"<span style='color: #6b7280; font-size: 13px;'>{pos[:50]}{'...' if len(pos) > 50 else ''}</span>", unsafe_allow_html=True)
        with col3:
            st.markdown(f"<span style='color: #9ca3af; font-size: 13px;'>{row['Applied on']}</span>", unsafe_allow_html=True)
        with col4:
            st.markdown(status_badge(row['Status']), unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# PAGE: APPLICATIONS
# ══════════════════════════════════════════════════════════
elif page == "📋 Applications":
    st.markdown("## Applications")

    col1, col2, col3 = st.columns([2, 1.5, 1])
    with col1:
        search = st.text_input("🔍 Search", placeholder="Company or position...")
    with col2:
        status_filter = st.selectbox("Status", ["All", "Applied", "Interviewed", "Offered", "Rejected", "Unknown"])
    with col3:
        limit = st.selectbox("Show", [50, 100, 200, "All"])

    filtered = df.copy()

    if search:
        mask = (
            filtered["Company"].str.contains(search, case=False, na=False) |
            filtered["Position"].str.contains(search, case=False, na=False)
        )
        filtered = filtered[mask]

    if status_filter != "All":
        filtered = filtered[filtered["Status"].str.contains(status_filter.lower(), case=False, na=False)]

    filtered = filtered.iloc[::-1]

    if limit != "All":
        filtered = filtered.head(int(limit))

    st.markdown(f"*Showing {len(filtered)} applications*")

    display_df = filtered[["Company", "Position", "Applied on", "Status"]].copy()
    display_df.columns = ["Company", "Position", "Applied", "Status"]

    st.dataframe(
        display_df,
        use_container_width=True,
        height=500,
        column_config={
            "Company": st.column_config.TextColumn("Company", width="medium"),
            "Position": st.column_config.TextColumn("Position", width="large"),
            "Applied": st.column_config.TextColumn("Applied", width="small"),
            "Status": st.column_config.TextColumn("Status", width="small"),
        }
    )


# ══════════════════════════════════════════════════════════
# PAGE: EMAIL SCANNER
# ══════════════════════════════════════════════════════════
elif page == "📧 Email Scanner":
    st.markdown("## Email Scanner")
    st.markdown("Scan your Gmail for job application updates and auto-update your Excel.")

    col1, col2 = st.columns([1, 2])
    with col1:
        num_emails = st.number_input("Emails to scan", min_value=5, max_value=200, value=50, step=5)

    if st.button("🔍 Scan emails now", use_container_width=False):
        with st.spinner("Connecting to Gmail and scanning emails..."):
            try:
                from tools import scan_emails_for_status
                result = scan_emails_for_status(max_results=num_emails)
                st.success(result)
                st.cache_data.clear()
            except Exception as e:
                st.error(f"Error: {e}")


# ══════════════════════════════════════════════════════════
# PAGE: COVER LETTER
# ══════════════════════════════════════════════════════════
elif page == "📝 Cover Letter":
    st.markdown("## Cover Letter Generator")

    # ── Session state init ───────────────────────────────────
    for _k, _v in [("cl_draft", None), ("cl_history", []), ("cl_jd", ""), ("cl_url", ""),
                     ("cl_download_ready", False), ("cl_download_buf", None), ("cl_download_name", "")]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # ── Step 1: Input (only shown when no draft exists) ──────
    if st.session_state["cl_draft"] is None:
        url = st.text_input("Job URL", placeholder="https://jobs.ashbyhq.com/...")
        jd_manual = st.text_area(
            "Or paste JD manually", height=200,
            placeholder="Paste the job description here if the URL can't be fetched automatically..."
        )

        if st.button("✨ Generate Cover Letter"):
            if not url and not jd_manual:
                st.warning("Please provide a URL or paste the JD.")
            else:
                with st.spinner("Generating cover letter..."):
                    try:
                        from nodes import _generate_cover_letter
                        from langchain_core.messages import HumanMessage
                        from tools import scrape_job_url

                        if url and not jd_manual:
                            jd = scrape_job_url(url)
                            if not jd or len(jd) < 200:
                                jd = jd_manual
                        else:
                            jd = jd_manual

                        if not jd:
                            st.error("Could not fetch JD. Please paste it manually.")
                        else:
                            result = _generate_cover_letter(
                                jd, url,
                                {"messages": [HumanMessage(content="generate cover letter")]}
                            )
                            draft = result["result"]
                            st.session_state["cl_draft"] = draft
                            st.session_state["cl_jd"] = jd
                            st.session_state["cl_url"] = url
                            st.session_state["cl_history"] = [("assistant", draft)]
                            st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")

    # ── Step 2: Draft editor + revision chat ─────────────────
    else:
        col_title, col_reset = st.columns([4, 1])
        with col_title:
            st.markdown("### Draft")
        with col_reset:
            if st.button("🔄 Start over"):
                for _k in ["cl_draft", "cl_history", "cl_jd", "cl_url"]:
                    st.session_state.pop(_k, None)
                st.rerun()

        # No key= so rerun always reflects the latest cl_draft value
        st.text_area(
            "You can edit directly or ask the agent to revise below:",
            value=st.session_state["cl_draft"],
            height=420,
        )

        st.markdown("---")

        # Revision history
        if len(st.session_state["cl_history"]) > 1:
            with st.expander("Revision history", expanded=False):
                for role, msg in st.session_state["cl_history"][1:]:
                    if role == "user":
                        st.markdown(
                            f"<div style='background:#f3f4f6;padding:8px 12px;border-radius:8px;"
                            f"margin:4px 0;font-size:13px;'>💬 {msg}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            "<div style='background:#eff6ff;padding:8px 12px;border-radius:8px;"
                            "margin:4px 0;font-size:13px;color:#1d4ed8;'>✏️ Draft updated</div>",
                            unsafe_allow_html=True,
                        )

        revision_input = st.text_input(
            "Ask the agent to revise:",
            placeholder="e.g. Make the opening more compelling, or shorten to 3 paragraphs",
        )

        col_revise, col_save = st.columns(2)

        with col_revise:
            if st.button("✏️ Revise", use_container_width=True):
                if not revision_input.strip():
                    st.warning("Please describe what you want to change.")
                else:
                    with st.spinner("Revising..."):
                        try:
                            from langchain_groq import ChatGroq
                            from langchain_core.messages import HumanMessage
                            llm = ChatGroq(model="meta-llama/llama-4-scout-17b-16e-instruct")
                            response = llm.invoke([HumanMessage(content=(
                                f"You are helping revise a cover letter.\n\n"
                                f"Current cover letter:\n{st.session_state['cl_draft']}\n\n"
                                f"Revision request: {revision_input}\n\n"
                                f"Return ONLY the revised cover letter, no explanation, no preamble."
                            ))])
                            new_draft = response.content.strip()
                            st.session_state["cl_history"].append(("user", revision_input))
                            st.session_state["cl_history"].append(("assistant", new_draft))
                            st.session_state["cl_draft"] = new_draft
                            st.rerun()
                        except Exception as e:
                            st.error(f"Revision failed: {e}")

        with col_save:
            if st.button("✅ Prepare download", use_container_width=True):
                with st.spinner("Preparing..."):
                    try:
                        from langchain_groq import ChatGroq
                        from langchain_core.messages import HumanMessage
                        from docx import Document
                        from docx.shared import Inches
                        import io

                        # Extract company and position from JD
                        llm = ChatGroq(model="meta-llama/llama-4-scout-17b-16e-instruct")
                        extract = llm.invoke([HumanMessage(content=(
                            f"Extract company name and job title from this job posting. "
                            f"Reply in this exact format:\ncompany: X\nposition: Y\n\n"
                            f"{st.session_state['cl_jd'][:2000]}"
                        ))])
                        company, position = "Unknown Company", "Unknown Position"
                        for line in extract.content.split("\n"):
                            if line.lower().startswith("company:"):
                                company = line.split(":", 1)[1].strip()
                            elif line.lower().startswith("position:"):
                                position = line.split(":", 1)[1].strip()

                        # Build docx in memory
                        doc = Document()
                        section = doc.sections[0]
                        section.top_margin    = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin   = Inches(0.5)
                        section.right_margin  = Inches(0.5)
                        doc.add_heading(f"Cover Letter — {position}", level=1)
                        doc.add_paragraph(f"Company: {company}")
                        doc.add_paragraph(f"Position: {position}")
                        doc.add_paragraph("")
                        for para in st.session_state["cl_draft"].split("\n"):
                            if para.strip():
                                doc.add_paragraph(para.strip())

                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)

                        filename = f"Cover_Letter_{position.replace(' ', '_')}.docx"
                        st.session_state["cl_download_buf"]  = buf.getvalue()
                        st.session_state["cl_download_name"] = filename
                        st.session_state["cl_download_ready"] = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"Failed: {e}")

        # Download button appears once the docx is ready
        if st.session_state.get("cl_download_ready"):
            st.download_button(
                label="⬇️ Download .docx",
                data=st.session_state["cl_download_buf"],
                file_name=st.session_state["cl_download_name"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            if st.button("🗑️ Done, clear draft", use_container_width=True):
                for _k in ["cl_draft", "cl_history", "cl_jd", "cl_url",
                           "cl_download_buf", "cl_download_name", "cl_download_ready"]:
                    st.session_state.pop(_k, None)
                st.rerun()

# ══════════════════════════════════════════════════════════
# PAGE: INTERVIEW PREP
# ══════════════════════════════════════════════════════════
elif page == "🎯 Interview Prep":
    st.markdown("## Interview Prep")

    url = st.text_input("Job URL", placeholder="https://...")
    jd_manual = st.text_area("Or paste JD manually", height=200, placeholder="Paste the job description here...")

    if st.button("🎯 Generate Interview Prep"):
        if not url and not jd_manual:
            st.warning("Please provide a URL or paste the JD.")
        else:
            with st.spinner("Generating interview prep guide..."):
                try:
                    from nodes import _generate_interview_prep
                    from langchain_core.messages import HumanMessage

                    if url and not jd_manual:
                        from tools import scrape_job_url
                        jd = scrape_job_url(url)
                        if not jd or len(jd) < 200:
                            jd = jd_manual
                    else:
                        jd = jd_manual

                    if not jd:
                        st.error("Could not fetch JD. Please paste it manually.")
                    else:
                        result = _generate_interview_prep(jd, url, {"messages": [HumanMessage(content="interview prep")]})
                        st.markdown(result["result"])

                except Exception as e:
                    st.error(f"Error: {e}")


# ══════════════════════════════════════════════════════════
# PAGE: JOB MATCH
# ══════════════════════════════════════════════════════════
elif page == "🔍 Job Match":
    st.markdown("## Job Match Scorer")

    tab1, tab2 = st.tabs(["Single job", "Batch score"])

    with tab1:
        company_input = st.text_input("Company name", placeholder="Caesars Entertainment")
        url = st.text_input("Or job URL", placeholder="https://...")
        jd_manual = st.text_area("Or paste JD manually", height=150, placeholder="Paste the job description here...")

        if st.button("🔍 Score my match"):
            with st.spinner("Analyzing job match..."):
                try:
                    from nodes import _generate_job_match
                    from tools import get_jd_from_excel, scrape_job_url
                    from langchain_core.messages import HumanMessage

                    jd, final_company, final_position, row_index = "", "", "", None

                    if company_input:
                        excel_data = get_jd_from_excel(company=company_input)
                        if excel_data and excel_data["jd"]:
                            jd = excel_data["jd"]
                            final_company = excel_data["company"]
                            final_position = excel_data["position"]
                            row_index = excel_data["row_index"]
                            if not url and excel_data["url"]:
                                url = excel_data["url"]

                    if not jd and url:
                        jd = scrape_job_url(url)

                    if not jd and jd_manual:
                        jd = jd_manual

                    if not jd:
                        st.error("Could not find JD. Please paste it manually.")
                    else:
                        result = _generate_job_match(jd, url, final_company, final_position, row_index, {"messages": [HumanMessage(content="job match")]})
                        st.markdown(result["result"])
                        st.cache_data.clear()

                except Exception as e:
                    st.error(f"Error: {e}")

    with tab2:
        limit = st.number_input("Score latest N jobs", min_value=1, max_value=20, value=5)

        if st.button("🔍 Batch score"):
            with st.spinner(f"Scoring {limit} jobs..."):
                try:
                    from tools import get_latest_jobs_with_jd, scrape_job_url
                    from nodes import _generate_job_match
                    from langchain_core.messages import HumanMessage

                    jobs = get_latest_jobs_with_jd(limit=int(limit))
                    results = []

                    progress = st.progress(0)
                    for idx, job in enumerate(jobs):
                        jd = str(job.get("JD") or "").strip()
                        url = str(job.get("Application Link") or "").strip()
                        job_company = str(job.get("Company") or "")
                        job_position = str(job.get("Position") or "")
                        row_index = job["row_index"]

                        if not jd and url:
                            jd = scrape_job_url(url)

                        if not jd:
                            results.append({"Company": job_company, "Position": job_position, "Score": "N/A", "Note": "No JD"})
                            continue

                        match_result = _generate_job_match(jd, url, job_company, job_position, row_index, {"messages": [HumanMessage(content="job match")]})

                        score = "N/A"
                        import re
                        for line in match_result["result"].split("\n"):
                            if "SCORE:" in line.upper():
                                m = re.search(r'(\d+)/10', line)
                                if m:
                                    score = f"{m.group(1)}/10"
                                    break

                        results.append({"Company": job_company, "Position": job_position[:50], "Score": score})
                        progress.progress((idx + 1) / len(jobs))

                    st.cache_data.clear()
                    st.dataframe(pd.DataFrame(results), use_container_width=True)

                except Exception as e:
                    st.error(f"Error: {e}")


# ══════════════════════════════════════════════════════════
# PAGE: CHAT
# ══════════════════════════════════════════════════════════
elif page == "💬 Chat":
    st.markdown("## Chat with your Agent")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "agent_state" not in st.session_state:
        st.session_state.agent_state = {
            "messages": [],
            "next_agent": "",
            "intent": "",
            "result": "",
            "pending_action": "",
            "pending_url": "",
            "pending_jd": ""
        }

    # Display chat history
    for msg in st.session_state.chat_history:
        if msg["role"] == "user":
            st.markdown(f'<div class="chat-msg-user">{msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-msg-agent">{msg["content"]}</div>', unsafe_allow_html=True)

    # Input
    with st.form("chat_form", clear_on_submit=True):
        col1, col2 = st.columns([5, 1])
        with col1:
            user_input = st.text_input("", placeholder="Ask anything... e.g. 'Show my latest 5 applications'", label_visibility="collapsed")
        with col2:
            submitted = st.form_submit_button("Send ↗", use_container_width=True)

    if submitted and user_input:
        from langchain_core.messages import HumanMessage

        st.session_state.chat_history.append({"role": "user", "content": user_input})
        st.session_state.agent_state["messages"].append(HumanMessage(content=user_input))

        with st.spinner("Thinking..."):
            try:
                from supervisor import supervisor
                result = supervisor.invoke({**st.session_state.agent_state})
                st.session_state.agent_state = {**st.session_state.agent_state, **result}
                response = result["messages"][-1].content
                st.session_state.chat_history.append({"role": "agent", "content": response})
                st.cache_data.clear()
            except Exception as e:
                st.session_state.chat_history.append({"role": "agent", "content": f"Error: {e}"})

        st.rerun()

    # Quick action buttons
    st.markdown("---")
    st.markdown("**Quick actions**")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("📋 Show latest 5 apps", use_container_width=True):
            st.session_state.chat_history.append({"role": "user", "content": "show me my latest 5 applications"})
            st.rerun()
    with col2:
        if st.button("📊 Job search analysis", use_container_width=True):
            st.session_state.chat_history.append({"role": "user", "content": "how is my job search going?"})
            st.rerun()
    with col3:
        if st.button("📧 Scan 20 emails", use_container_width=True):
            st.session_state.chat_history.append({"role": "user", "content": "scan my last 20 emails for job updates"})
            st.rerun()
    with col4:
        if st.button("🗑️ Clear chat", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.agent_state = {
                "messages": [], "next_agent": "", "intent": "",
                "result": "", "pending_action": "", "pending_url": "", "pending_jd": ""
            }
            st.rerun()


# ══════════════════════════════════════════════════════════
# PAGE: SETUP
# ══════════════════════════════════════════════════════════
elif page == "⚙️ Setup":
    from onboarding import render_setup_page
    render_setup_page()
